"""Upload handling — CV documents and profile photos (local disk or Firebase Storage)."""

from __future__ import annotations

import io
import os
import shutil
from typing import Optional, Tuple

from backend.config import DATA_DIR
from backend.firebase_app import is_enabled
from backend.models import CVContent, CVDocument, WritingTone
from backend import ai_service

UPLOADS_DIR = os.path.join(str(DATA_DIR), "uploads")
MAX_CV_BYTES = 10 * 1024 * 1024
MAX_PHOTO_BYTES = 5 * 1024 * 1024

CV_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
PHOTO_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}


def _cv_dir(cv_id: str) -> str:
    path = os.path.join(UPLOADS_DIR, cv_id)
    os.makedirs(path, exist_ok=True)
    return path


def _ext(filename: str) -> str:
    return os.path.splitext(filename or "")[1].lower()


def _extract_pdf_pypdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    try:
        reader = PdfReader(io.BytesIO(data))
        parts = [(page.extract_text() or "").strip() for page in reader.pages]
        return "\n\n".join(p for p in parts if p)
    except Exception:
        return ""


def _extract_pdf_pdfplumber(data: bytes) -> str:
    """More robust extractor (pdfminer-based) for PDFs pypdf can't read well."""
    try:
        import pdfplumber
    except ImportError:
        return ""
    try:
        parts = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                parts.append((page.extract_text() or "").strip())
        return "\n\n".join(p for p in parts if p)
    except Exception:
        return ""


def _extract_pdf_pdfminer(data: bytes) -> str:
    try:
        from pdfminer.high_level import extract_text
    except ImportError:
        return ""
    try:
        return (extract_text(io.BytesIO(data)) or "").strip()
    except Exception:
        return ""


def extract_cv_text(filename: str, data: bytes) -> str:
    ext = _ext(filename)
    if ext == ".pdf":
        # Try multiple engines and keep the best (longest) extraction.
        best = ""
        for engine in (_extract_pdf_pypdf, _extract_pdf_pdfplumber, _extract_pdf_pdfminer):
            text = engine(data)
            if len(text) > len(best):
                best = text
            if len(best) >= 200:
                break
        return best

    if ext == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required for DOCX upload") from exc
        doc = Document(io.BytesIO(data))
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)

    if ext in {".txt", ".md"}:
        return data.decode("utf-8", errors="ignore")

    raise ValueError(f"Unsupported CV file type: {ext or 'unknown'}. Use PDF, DOCX, or TXT.")


def save_profile_photo(user_id: str, cv_id: str, filename: str, data: bytes) -> str:
    if is_enabled():
        from backend import firebase_upload
        if len(data) > MAX_PHOTO_BYTES:
            raise ValueError("Profile photo must be under 5 MB.")
        return firebase_upload.save_profile_photo(user_id, cv_id, filename, data)

    del user_id
    ext = _ext(filename)
    if ext not in PHOTO_EXTENSIONS:
        raise ValueError("Profile photo must be JPG, PNG, or WebP.")
    if len(data) > MAX_PHOTO_BYTES:
        raise ValueError("Profile photo must be under 5 MB.")

    dest_dir = _cv_dir(cv_id)
    for old in os.listdir(dest_dir):
        if old.startswith("profile."):
            try:
                os.remove(os.path.join(dest_dir, old))
            except OSError:
                pass

    dest = os.path.join(dest_dir, f"profile{ext}")
    with open(dest, "wb") as f:
        f.write(data)
    return f"/api/cvs/{cv_id}/photo"


def get_photo_bytes(user_id: str, cv_id: str) -> Optional[Tuple[bytes, str]]:
    if is_enabled():
        from backend import firebase_upload
        return firebase_upload.get_photo_bytes(user_id, cv_id)

    del user_id
    dest_dir = os.path.join(UPLOADS_DIR, cv_id)
    if not os.path.isdir(dest_dir):
        return None
    for name in sorted(os.listdir(dest_dir)):
        if name.startswith("profile.") and _ext(name) in PHOTO_EXTENSIONS:
            path = os.path.join(dest_dir, name)
            ext = _ext(name)
            media = "image/jpeg"
            if ext == ".png":
                media = "image/png"
            elif ext == ".webp":
                media = "image/webp"
            with open(path, "rb") as f:
                return f.read(), media
    return None


def delete_uploads(user_id: str, cv_id: str) -> None:
    if is_enabled():
        from backend import firebase_upload
        firebase_upload.delete_uploads(user_id, cv_id)
        return

    del user_id
    path = os.path.join(UPLOADS_DIR, cv_id)
    if os.path.isdir(path):
        shutil.rmtree(path, ignore_errors=True)


def copy_uploads(user_id: str, source_id: str, dest_id: str) -> None:
    if is_enabled():
        from backend import firebase_upload
        firebase_upload.copy_uploads(user_id, source_id, dest_id)
        return

    del user_id
    src = os.path.join(UPLOADS_DIR, source_id)
    if not os.path.isdir(src):
        return
    dst = _cv_dir(dest_id)
    for name in os.listdir(src):
        shutil.copy2(os.path.join(src, name), os.path.join(dst, name))


def process_cv_file(
    doc: CVDocument,
    filename: str,
    data: bytes,
    tone: WritingTone = WritingTone.PROFESSIONAL,
) -> Tuple[CVDocument, str]:
    if len(data) > MAX_CV_BYTES:
        raise ValueError("CV file must be under 10 MB.")

    ext = _ext(filename)
    if ext not in CV_EXTENSIONS:
        raise ValueError("Upload PDF, DOCX, or TXT resume file.")

    raw_text = extract_cv_text(filename, data).strip()
    if len(raw_text) < 30:
        raise ValueError(
            "Could not read text from this file. If it's a scanned or image-based PDF, "
            "the text can't be extracted — please upload a text-based PDF or DOCX, "
            "or paste your CV details into the chat."
        )

    result = ai_service.process_cv_upload(raw_text, doc.content, tone)
    if not result.success:
        raise RuntimeError(result.message or "AI could not parse your CV file.")

    updated_content = CVContent(**result.data["content"])
    new_doc = doc.model_copy(deep=True)
    new_doc.content = updated_content
    if updated_content.full_name:
        new_doc.name = f"{updated_content.full_name} CV"

    return new_doc, (
        f"I've imported your CV from **{filename}** ({len(raw_text):,} characters extracted). "
        "Review the preview — you can edit anything via chat."
    )
