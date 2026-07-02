"""Export CV to PDF and DOCX."""

from __future__ import annotations

import io
import os
import tempfile
from typing import Tuple

from backend.models import CVDocument
from backend.templates import get_template


def _section_title(name: str) -> str:
    return name.replace("_", " ").title()


def _build_plain_text(doc: CVDocument) -> str:
    c = doc.content
    lines = [
        c.full_name.upper(),
        c.job_title,
        " | ".join(filter(None, [
            c.contact.email, c.contact.phone, c.contact.location,
            c.contact.linkedin, c.contact.website,
        ])),
        "",
    ]

    visibility = c.section_visibility.model_dump()
    for section in c.section_order:
        if not visibility.get(section, True):
            continue

        if section == "summary" and c.summary:
            lines += [_section_title(section), c.summary, ""]
        elif section == "experience" and c.experience:
            lines.append(_section_title(section))
            for exp in c.experience:
                dates = f"{exp.start_date} – {exp.end_date or 'Present'}" if exp.start_date else ""
                lines.append(f"{exp.role} | {exp.company} | {dates}")
                for b in exp.bullets:
                    lines.append(f"  • {b}")
                lines.append("")
        elif section == "education" and c.education:
            lines.append(_section_title(section))
            for edu in c.education:
                lines.append(f"{edu.degree} in {edu.field} — {edu.institution} ({edu.start_date}–{edu.end_date})")
            lines.append("")
        elif section == "projects" and c.projects:
            lines.append(_section_title(section))
            for p in c.projects:
                lines.append(f"{p.name}: {p.description}")
                for b in p.bullets:
                    lines.append(f"  • {b}")
            lines.append("")
        elif section == "skills" and c.skills:
            lines += [_section_title(section), ", ".join(c.skills), ""]
        elif section == "certifications" and c.certifications:
            lines += [_section_title(section), "\n".join(f"• {x}" for x in c.certifications), ""]
        elif section == "languages" and c.languages:
            lines += [_section_title(section), ", ".join(c.languages), ""]
        elif section == "awards" and c.awards:
            lines += [_section_title(section), "\n".join(f"• {x}" for x in c.awards), ""]

    return "\n".join(lines).strip() + "\n"


def export_docx(doc: CVDocument) -> Tuple[bytes, str]:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX export. pip install python-docx") from exc

    template = get_template(doc.template_id)
    c = doc.content
    document = Document()

    title = document.add_heading(c.full_name or doc.name, level=0)
    if c.job_title:
        sub = document.add_paragraph(c.job_title)
        sub.runs[0].font.size = Pt(12)
        sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    contact_bits = [c.contact.email, c.contact.phone, c.contact.location, c.contact.linkedin]
    if any(contact_bits):
        document.add_paragraph(" | ".join(x for x in contact_bits if x))

    visibility = c.section_visibility.model_dump()
    for section in c.section_order:
        if not visibility.get(section, True):
            continue

        if section == "summary" and c.summary:
            document.add_heading("Summary", level=1)
            document.add_paragraph(c.summary)
        elif section == "experience" and c.experience:
            document.add_heading("Experience", level=1)
            for exp in c.experience:
                p = document.add_paragraph()
                p.add_run(f"{exp.role} — {exp.company}").bold = True
                for b in exp.bullets:
                    document.add_paragraph(b, style="List Bullet")
        elif section == "education" and c.education:
            document.add_heading("Education", level=1)
            for edu in c.education:
                document.add_paragraph(f"{edu.degree}, {edu.field} — {edu.institution}")
        elif section == "projects" and c.projects:
            document.add_heading("Projects", level=1)
            for proj in c.projects:
                document.add_paragraph(proj.name).runs[0].bold = True
                if proj.description:
                    document.add_paragraph(proj.description)
        elif section == "skills" and c.skills:
            document.add_heading("Skills", level=1)
            document.add_paragraph(", ".join(c.skills))

    buf = io.BytesIO()
    document.save(buf)
    filename = f"{(c.full_name or doc.name).replace(' ', '_')}.docx"
    return buf.getvalue(), filename


def export_pdf(doc: CVDocument) -> Tuple[bytes, str]:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError("reportlab is required for PDF export. pip install reportlab") from exc

    template = get_template(doc.template_id)
    buf = io.BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=letter, leftMargin=0.75 * inch, rightMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    accent = template.get("preview_color", "#1d4ed8")

    title_style = ParagraphStyle(
        "CVTitle",
        parent=styles["Heading1"],
        fontSize=20,
        spaceAfter=4,
        textColor=accent,
    )
    section_style = ParagraphStyle(
        "CVSection",
        parent=styles["Heading2"],
        fontSize=12,
        spaceBefore=10,
        spaceAfter=4,
        textColor=accent,
    )
    body_style = styles["BodyText"]

    story = []
    c = doc.content

    story.append(Paragraph(c.full_name or doc.name, title_style))
    if c.job_title:
        story.append(Paragraph(c.job_title, body_style))
    contact = " | ".join(filter(None, [c.contact.email, c.contact.phone, c.contact.location]))
    if contact:
        story.append(Paragraph(contact, body_style))
    story.append(Spacer(1, 8))

    text = _build_plain_text(doc)
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        first_line = block.split("\n")[0]
        if first_line in {_section_title(s) for s in c.section_order}:
            story.append(Paragraph(first_line, section_style))
            rest = "\n".join(block.split("\n")[1:]).strip()
            if rest:
                story.append(Paragraph(rest.replace("\n", "<br/>"), body_style))
        else:
            story.append(Paragraph(block.replace("\n", "<br/>"), body_style))

    pdf.build(story)
    filename = f"{(c.full_name or doc.name).replace(' ', '_')}.pdf"
    return buf.getvalue(), filename
